

from docx import Document
from docx.shared import Pt

def txt_to_docx_arial10(txt_path, docx_path):
    document = Document()
    with open(txt_path, 'r', encoding='utf-8') as txt_file:
        content = txt_file.read()
        paragraph = document.add_paragraph()
        run = paragraph.add_run(content)
        font = run.font
        font.name = 'Arial'
        font.size = Pt(10)
    document.save(docx_path)

# Uso: 
txt_to_docx_arial10('documento.txt', 'examen.docx')



